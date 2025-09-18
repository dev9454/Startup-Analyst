# tools/loaders.py
from __future__ import annotations
import os, re, pathlib, html
from typing import List
import requests

from langchain.schema import Document  # ok with current LC
# If you use langchain-core >=0.2: from langchain_core.documents import Document

from tools.docai_ocr import docai_ocr_pdf_bytes

# Optional HTML -> text helpers (pure stdlib fallback)
def _strip_html(raw: str) -> str:
    # very light cleanup; you can replace with 'trafilatura' if you like
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", raw)
    text = re.sub(r"(?is)<br\s*/?>", "\n", text)
    text = re.sub(r"(?is)</p>", "\n\n", text)
    text = re.sub(r"(?is)<.*?>", " ", text)
    text = html.unescape(text)
    # collapse spaces
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    return text.strip()

def _is_url(s: str) -> bool:
    return s.lower().startswith("http://") or s.lower().startswith("https://")

def load_pdf_docai(path_or_bytes: str | bytes) -> List[Document]:
    if isinstance(path_or_bytes, (str, os.PathLike)):
        with open(path_or_bytes, "rb") as f:
            pdf_bytes = f.read()
        source = str(path_or_bytes)
    else:
        pdf_bytes = path_or_bytes
        source = "buffer.pdf"

    pages = docai_ocr_pdf_bytes(pdf_bytes)  # list[(text, conf)]
    docs: List[Document] = []
    for i, (txt, _conf) in enumerate(pages, start=1):
        if not txt or not txt.strip():
            continue
        docs.append(
            Document(
                page_content=txt.strip(),
                metadata={
                    "source": source,
                    "page": i,
                    "loader": "docai_ocr",
                },
            )
        )
    return docs

def load_docx(path: str) -> List[Document]:
    try:
        import docx  # python-docx
    except Exception:
        # fallback: very naive read
        return [Document(page_content=open(path, "rb").read().decode("utf-8", errors="ignore"),
                         metadata={"source": path, "loader": "raw-bytes"})]
    doc = docx.Document(path)
    paras = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n".join(paras).strip()
    return [Document(page_content=text, metadata={"source": path, "loader": "docx"})] if text else []

def load_url(url: str) -> List[Document]:
    headers = {"User-Agent": os.getenv("USER_AGENT", "StartupAnalyst/1.0")}
    r = requests.get(url, headers=headers, timeout=30)
    r.raise_for_status()
    txt = _strip_html(r.text)
    # Keep a sane chunk size ~4000 chars to help embedding
    chunks: List[Document] = []
    CHUNK = 4000
    for i in range(0, len(txt), CHUNK):
        piece = txt[i:i+CHUNK].strip()
        if piece:
            chunks.append(
                Document(
                    page_content=piece,
                    metadata={"source": url, "chunk": i//CHUNK, "loader": "url"},
                )
            )
    return chunks

def load_text_file(path: str) -> List[Document]:
    data = open(path, "r", encoding="utf-8", errors="ignore").read()
    return [Document(page_content=data, metadata={"source": path, "loader": "text"})]

def load_many(inputs: List[str]) -> List[Document]:
    out: List[Document] = []
    for it in inputs:
        it = str(it).strip().strip('"').strip("'")
        if not it:
            continue

        if _is_url(it):
            try:
                out.extend(load_url(it))
            except Exception as e:
                out.append(Document(page_content=f"[url_error] {e}", metadata={"source": it, "loader": "url"}))
            continue

        p = pathlib.Path(it)
        if not p.exists():
            out.append(Document(page_content=f"[missing] {it}", metadata={"source": it, "loader": "missing"}))
            continue

        ext = p.suffix.lower()
        try:
            if ext == ".pdf":
                # ALWAYS OCR for PDFs
                out.extend(load_pdf_docai(str(p)))
            elif ext in (".docx", ".doc"):
                out.extend(load_docx(str(p)))
            elif ext in (".txt", ".md", ".csv"):
                out.extend(load_text_file(str(p)))
            else:
                # Try OCR for unknown binaries? Keep simple: read as text
                out.extend(load_text_file(str(p)))
        except Exception as e:
            out.append(Document(page_content=f"[load_error] {e}", metadata={"source": str(p), "loader": "error"}))

    return out
